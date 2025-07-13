import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import re

class WordDocumentGenerator:
    """
    Utility class to generate Word documents for SOPs and Batch Records
    with proper formatting and prefilled sections
    """
    
    def __init__(self, output_dir="generated_docs"):
        self.output_dir = output_dir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
    
    def generate_sop_document(self, content, title=None, doc_id=None, template_type=None):
        """
        Generate a Word document for an SOP with proper formatting
        
        Args:
            content (str): The content of the SOP
            title (str): The title of the SOP
            doc_id (str): Document ID for tracking
            template_type (str): Type of template to use (e.g., "NK_cell_thawing")
            
        Returns:
            str: Path to the generated document
        """
        # Create a new document
        doc = Document()
        
        # Set up document styles
        self._setup_document_styles(doc)
        
        # Add header with logo placeholder
        self._add_header(doc, "Standard Operating Procedure (SOP)")
        
        # Extract title from content if not provided
        if not title:
            title_match = re.search(r"Title:\s*(.*?)(?:\n|$)", content)
            if title_match:
                title = title_match.group(1).strip()
            else:
                title = "Standard Operating Procedure"
        
        # Add title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add document info table
        self._add_document_info_table(doc, doc_id)
        
        # Process content based on template type
        if template_type and template_type.lower() == "nk_cell_thawing":
            self._add_nk_cell_thawing_template(doc, content)
        else:
            # Process general content
            self._add_general_content(doc, content)
            
        # Add footer
        self._add_footer(doc)
        
        # Save the document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')
        filename = f"{safe_title}_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        doc.save(filepath)
        return filepath
    
    def _setup_document_styles(self, doc):
        """Set up document styles"""
        # Modify the Normal style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Create a Heading 1 style
        if 'Heading 1' not in doc.styles:
            style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles['Heading 1']
        
        font = style.font
        font.name = 'Arial'
        font.size = Pt(14)
        font.bold = True
        
        # Create a Heading 2 style
        if 'Heading 2' not in doc.styles:
            style = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles['Heading 2']
        
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = True
    
    def _add_header(self, doc, doc_type):
        """Add a header to the document"""
        section = doc.sections[0]
        header = section.header
        
        # Add a table for the header
        table = header.add_table(1, 2, Inches(6))
        
        # Company name/logo placeholder
        company_cell = table.cell(0, 0)
        company_paragraph = company_cell.paragraphs[0]
        company_run = company_paragraph.add_run("COMPANY LOGO")
        company_run.bold = True
        company_run.font.size = Pt(12)
        
        # Document type
        type_cell = table.cell(0, 1)
        type_paragraph = type_cell.paragraphs[0]
        type_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        type_run = type_paragraph.add_run(doc_type)
        type_run.bold = True
        type_run.font.size = Pt(12)
        
        # Add a line after the header
        doc.add_paragraph().add_run().add_break()
    
    def _add_document_info_table(self, doc, doc_id):
        """Add document information table"""
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # Document ID
        table.cell(0, 0).text = "Document ID:"
        table.cell(0, 1).text = doc_id if doc_id else "SOP-" + datetime.now().strftime("%Y%m%d")
        
        # Effective Date
        table.cell(1, 0).text = "Effective Date:"
        table.cell(1, 1).text = datetime.now().strftime("%Y-%m-%d")
        
        # Revision Number
        table.cell(2, 0).text = "Revision Number:"
        table.cell(2, 1).text = "1.0"
        
        # Approval Status
        table.cell(3, 0).text = "Approval Status:"
        table.cell(3, 1).text = "Draft"
        
        # Format the table
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        # Add space after the table
        doc.add_paragraph()
    
    def _add_general_content(self, doc, content):
        """Add general content with section parsing"""
        # Split content by sections
        sections = re.split(r'\n\s*(?=\d+\.\s+[A-Z])', content)
        
        # Process each section
        for section in sections:
            # Check if this is the first section (might contain preamble)
            if sections.index(section) == 0 and not re.match(r'^\d+\.\s+[A-Z]', section.strip()):
                # Add preamble as regular paragraphs
                for line in section.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
                continue
            
            # Try to extract section title
            title_match = re.match(r'(\d+\.\s+[^\n]+)', section)
            if title_match:
                title = title_match.group(1).strip()
                # Add section title as Heading 1
                doc.add_paragraph(title, style='Heading 1')
                
                # Add section content
                content = section[len(title):].strip()
                if content:
                    # Check for subsections
                    subsections = re.split(r'\n\s*(?=\d+\.\d+\s+[A-Z])', content)
                    
                    for subsection in subsections:
                        # Check if this is a proper subsection
                        subtitle_match = re.match(r'(\d+\.\d+\s+[^\n]+)', subsection)
                        if subtitle_match:
                            subtitle = subtitle_match.group(1).strip()
                            # Add subsection title as Heading 2
                            doc.add_paragraph(subtitle, style='Heading 2')
                            
                            # Add subsection content
                            subcontent = subsection[len(subtitle):].strip()
                            if subcontent:
                                for line in subcontent.split('\n'):
                                    if line.strip():
                                        doc.add_paragraph(line.strip())
                        else:
                            # No proper subsection format, add as regular paragraph
                            for line in subsection.split('\n'):
                                if line.strip():
                                    doc.add_paragraph(line.strip())
            else:
                # No proper section format, add as regular paragraph
                for line in section.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
    
    def _add_nk_cell_thawing_template(self, doc, content):
        """Add NK cell thawing specific template with prefilled sections"""
        # Add purpose section
        doc.add_paragraph("1. PURPOSE", style='Heading 1')
        purpose_para = doc.add_paragraph()
        purpose_para.add_run("This Standard Operating Procedure (SOP) describes the process for thawing Natural Killer (NK) cells while maintaining cell viability and functionality for downstream applications.")
        
        # Add scope section
        doc.add_paragraph("2. SCOPE", style='Heading 1')
        scope_para = doc.add_paragraph()
        scope_para.add_run("This procedure applies to all laboratory personnel involved in the handling and processing of cryopreserved NK cells.")
        
        # Add responsibilities section
        doc.add_paragraph("3. RESPONSIBILITIES", style='Heading 1')
        resp_para = doc.add_paragraph()
        resp_para.add_run("It is the responsibility of all trained laboratory personnel to follow this SOP when thawing NK cells. The Laboratory Supervisor is responsible for ensuring that personnel are properly trained on this procedure.")
        
        # Add materials section
        doc.add_paragraph("4. MATERIALS AND EQUIPMENT", style='Heading 1')
        materials = [
            "Personal Protective Equipment (PPE): lab coat, gloves, safety glasses",
            "Water bath set to 37°C",
            "Timer",
            "70% ethanol spray",
            "Sterile serological pipettes (5 mL, 10 mL, 25 mL)",
            "Pipette controller",
            "Centrifuge tubes (15 mL, 50 mL)",
            "Complete culture medium (pre-warmed to 37°C)",
            "Centrifuge",
            "Biosafety cabinet (BSC)",
            "Cell counting equipment (hemocytometer or automated cell counter)",
            "Trypan blue solution (0.4%)",
            "Cryovial containing frozen NK cells"
        ]
        
        for item in materials:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)
        
        # Add procedure section
        doc.add_paragraph("5. PROCEDURE", style='Heading 1')
        
        # 5.1 Preparation
        doc.add_paragraph("5.1 Preparation", style='Heading 2')
        prep_steps = [
            "Ensure all required materials and equipment are available.",
            "Turn on the biosafety cabinet and allow it to run for at least 15 minutes before use.",
            "Set the water bath to 37°C and verify the temperature with a thermometer.",
            "Pre-warm complete culture medium to 37°C.",
            "Label all tubes clearly with the sample information."
        ]
        
        for i, step in enumerate(prep_steps, 1):
            p = doc.add_paragraph()
            p.add_run(f"5.1.{i} ").bold = True
            p.add_run(step)
        
        # 5.2 Thawing Procedure
        doc.add_paragraph("5.2 Thawing Procedure", style='Heading 2')
        thaw_steps = [
            "Remove the cryovial containing NK cells from liquid nitrogen storage and immediately place it into a container with dry ice or a portable LN2 container.",
            "Transport the cryovial to the water bath area.",
            "Partially submerge the cryovial in the 37°C water bath, ensuring the cap remains above the water level to prevent contamination.",
            "Gently swirl the vial in the water bath until only a small ice crystal remains (approximately 1-2 minutes).",
            "Spray the outside of the vial with 70% ethanol and transfer it to the biosafety cabinet.",
            "Using a 5 mL serological pipette, slowly transfer the cell suspension to a 15 mL centrifuge tube.",
            "Add pre-warmed complete culture medium dropwise to the cells, starting with 1 mL over the first minute while gently swirling the tube.",
            "Continue adding medium slowly, 1 mL at a time with gentle mixing, until reaching 5 mL total volume.",
            "Add the remaining medium to reach 10 mL total volume.",
            "Centrifuge the cell suspension at 300 × g for 5 minutes at room temperature.",
            "Carefully aspirate and discard the supernatant without disturbing the cell pellet.",
            "Gently resuspend the cell pellet in 5-10 mL of fresh pre-warmed complete culture medium."
        ]
        
        for i, step in enumerate(thaw_steps, 1):
            p = doc.add_paragraph()
            p.add_run(f"5.2.{i} ").bold = True
            p.add_run(step)
        
        # 5.3 Cell Counting and Viability Assessment
        doc.add_paragraph("5.3 Cell Counting and Viability Assessment", style='Heading 2')
        count_steps = [
            "Mix 10 μL of cell suspension with 10 μL of 0.4% trypan blue solution.",
            "Load 10 μL of the mixture onto a hemocytometer or use an automated cell counter according to the manufacturer's instructions.",
            "Count the number of viable (unstained) and non-viable (blue-stained) cells.",
            "Calculate the cell concentration and viability percentage.",
            "Record the cell count and viability in the laboratory notebook."
        ]
        
        for i, step in enumerate(count_steps, 1):
            p = doc.add_paragraph()
            p.add_run(f"5.3.{i} ").bold = True
            p.add_run(step)
        
        # 5.4 Post-Thaw Culture
        doc.add_paragraph("5.4 Post-Thaw Culture", style='Heading 2')
        culture_steps = [
            "Adjust the cell concentration to the required density for your specific application (typically 0.5-1 × 10^6 cells/mL).",
            "Transfer the cells to an appropriate culture vessel.",
            "Incubate the cells at 37°C, 5% CO2 in a humidified incubator.",
            "Monitor cell recovery and proliferation after 24 hours."
        ]
        
        for i, step in enumerate(culture_steps, 1):
            p = doc.add_paragraph()
            p.add_run(f"5.4.{i} ").bold = True
            p.add_run(step)
        
        # Add quality control section
        doc.add_paragraph("6. QUALITY CONTROL", style='Heading 1')
        qc_para = doc.add_paragraph()
        qc_para.add_run("Cell viability should be ≥ 70% post-thaw. If viability is consistently below this threshold, review and optimize the freezing and thawing procedures.")
        
        # Add references section
        doc.add_paragraph("7. REFERENCES", style='Heading 1')
        ref_items = [
            "Current Good Manufacturing Practice (cGMP) guidelines",
            "Manufacturer's instructions for equipment and materials used",
            "Laboratory safety manual"
        ]
        
        for item in ref_items:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)
        
        # Add revision history section
        doc.add_paragraph("8. REVISION HISTORY", style='Heading 1')
        
        # Create revision history table
        rev_table = doc.add_table(rows=2, cols=4)
        rev_table.style = 'Table Grid'
        
        # Add headers
        headers = ["Revision Number", "Effective Date", "Description of Changes", "Author"]
        for i, header in enumerate(headers):
            rev_table.cell(0, i).text = header
            for paragraph in rev_table.cell(0, i).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add initial revision
        rev_data = ["1.0", datetime.now().strftime("%Y-%m-%d"), "Initial release", ""]
        for i, data in enumerate(rev_data):
            rev_table.cell(1, i).text = data
    
    def _add_footer(self, doc):
        """Add a footer to the document"""
        section = doc.sections[0]
        footer = section.footer
        
        # Add a paragraph for the footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number
        footer_run = footer_paragraph.add_run("Page ")
        footer_run.font.size = Pt(8)
        
        # Add date and confidentiality statement
        date_paragraph = footer.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d')} | CONFIDENTIAL")
        date_run.font.size = Pt(8)
