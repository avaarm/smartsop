"""Rich DOCX structure extractor.

Goes beyond python-docx's built-in helpers to pull out the formatting
details that matter for faithful reproduction:

  * Page setup (orientation, margins, size in DXA)
  * Headers and footers (first-page + default) as text
  * Per-role font histograms (body, Heading 1, Heading 2, ...) so
    downstream analysis can state "body is Calibri 11pt" with
    confidence instead of guessing from a global top font.
  * Rich table extraction:
      - Column widths (DXA twips)
      - Per-cell shading (background fill colors, e.g. ``BFBFBF``)
      - Cell borders (top/left/bottom/right: style, size, color)
      - Horizontal merges (``w:gridSpan``)
      - Vertical merges (``w:vMerge``)
      - Header-row flag (``w:tblHeader``)
      - Per-cell text (first 200 chars)
  * Numbering scheme (already in parser, mirrored here for cohesion).

The output is pure JSON-serializable dicts so the ProtocolAnalyzer
can feed it to the LLM, and the frontend can render it verbatim in
the Formatting knowledge card.
"""

from __future__ import annotations

import logging
import re
from collections import Counter, defaultdict
from typing import Optional

logger = logging.getLogger(__name__)

# OOXML namespaces — every qualified tag under w:// lives here.
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def _qn(tag: str) -> str:
    """Return a `{ns}localname` Clark-notation tag for direct lxml lookups."""
    return f"{{{W_NS}}}{tag}"


def _attr(elem, tag: str) -> Optional[str]:
    """Read a ``w:*`` attribute off an element, returning None if absent."""
    return elem.get(_qn(tag)) if elem is not None else None


# ── Top-level extractor ────────────────────────────────────────────

def extract_docx_structure(file_path: str) -> dict:
    """Return a rich structure dict for ``file_path``.

    Shape:
        {
            "page": {"orientation", "width_dxa", "height_dxa",
                     "margins_dxa": {top, right, bottom, left}},
            "headers": {"default": "...", "first_page": "..."},
            "footers": {"default": "...", "first_page": "..."},
            "font_roles": {
                "body": {"name", "size_pt", "bold", "italic", "count"},
                "Heading 1": {...}, ...
            },
            "shading_palette": [{"color": "BFBFBF", "count": 12}, ...],
            "tables": [
                {
                    "index", "rows", "cols",
                    "col_widths_dxa": [int, ...],
                    "has_header_row": bool,
                    "cells": [[{"text", "shading", "gridSpan",
                                "vMerge", "borders"}, ...], ...]
                }
            ],
            "numbering_scheme": "X.X.X" | "X" | "mixed" | "none",
        }
    """
    from docx import Document

    doc = Document(file_path)

    # Theme fonts are referenced by the docx-generic name "Default" in
    # python-docx; resolve them from theme1.xml so we can report the
    # actual font (e.g. Calibri / Calibri Light) instead of "Default".
    theme_fonts = _extract_theme_fonts(doc)

    return {
        "page": _extract_page_setup(doc),
        "headers": _extract_headers(doc),
        "footers": _extract_footers(doc),
        "font_roles": _extract_font_roles(doc, theme_fonts),
        "paragraph_rhythm": _extract_paragraph_rhythm(doc),
        "numbering": _extract_numbering_definitions(doc),
        "tables": [_extract_table(t, idx) for idx, t in enumerate(doc.tables)],
        "shading_palette": _extract_shading_palette(doc),
        "theme_fonts": theme_fonts,
    }


# ── Page setup ─────────────────────────────────────────────────────

def _extract_page_setup(doc) -> dict:
    """Read the first section's ``<w:sectPr>`` for page size + margins."""
    try:
        sect = doc.sections[0]
    except IndexError:
        return {}

    sect_pr = sect._sectPr
    pg_sz = sect_pr.find(_qn("pgSz"))
    pg_mar = sect_pr.find(_qn("pgMar"))

    orientation = "portrait"
    width_dxa = None
    height_dxa = None
    if pg_sz is not None:
        width_dxa = int(_attr(pg_sz, "w") or 0) or None
        height_dxa = int(_attr(pg_sz, "h") or 0) or None
        orient = _attr(pg_sz, "orient")
        if orient:
            orientation = orient
        elif width_dxa and height_dxa and width_dxa > height_dxa:
            orientation = "landscape"

    margins = {}
    if pg_mar is not None:
        for side in ("top", "right", "bottom", "left"):
            val = _attr(pg_mar, side)
            if val is not None:
                try:
                    margins[f"{side}_dxa"] = int(val)
                except ValueError:
                    pass

    return {
        "orientation": orientation,
        "width_dxa": width_dxa,
        "height_dxa": height_dxa,
        "margins_dxa": margins,
        "width_inches": round(width_dxa / 1440, 2) if width_dxa else None,
        "height_inches": round(height_dxa / 1440, 2) if height_dxa else None,
    }


# ── Headers and footers ────────────────────────────────────────────

def _extract_headers(doc) -> dict:
    """Return header text keyed by kind (default / first_page / even_page)."""
    out = {}
    for sect in doc.sections:
        for kind, header in (
            ("default", sect.header),
            ("first_page", sect.first_page_header),
            ("even_page", sect.even_page_header),
        ):
            try:
                text = "\n".join(p.text for p in header.paragraphs if p.text.strip())
            except Exception:
                text = ""
            if text and kind not in out:
                out[kind] = text[:1000]
    return out


def _extract_footers(doc) -> dict:
    out = {}
    for sect in doc.sections:
        for kind, footer in (
            ("default", sect.footer),
            ("first_page", sect.first_page_footer),
            ("even_page", sect.even_page_footer),
        ):
            try:
                text = "\n".join(p.text for p in footer.paragraphs if p.text.strip())
            except Exception:
                text = ""
            if text and kind not in out:
                out[kind] = text[:1000]
    return out


# ── Fonts by role ──────────────────────────────────────────────────

# ── Theme fonts ────────────────────────────────────────────────────

_THEME_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def _extract_theme_fonts(doc) -> dict:
    """Resolve the document's theme1.xml major/minor Latin fonts.

    Word docs almost never hard-code "Calibri" — they reference the theme
    font via ``+mj-lt`` (major heading) or ``+mn-lt`` (minor body). Without
    resolving these, every font lookup returns "Default" which is useless
    for a downstream generator. Returns a dict:

        {"major": "Calibri Light", "minor": "Calibri"}
    """
    try:
        theme_part = doc.part.package.part_related_by.__self__  # noqa - not used
    except Exception:
        pass

    try:
        # python-docx exposes the theme through the document part's
        # related-parts collection. Walk them for the theme XML.
        for rel in doc.part.rels.values():
            if "theme" in rel.reltype:
                theme_xml = rel.target_part.blob
                return _parse_theme_fonts(theme_xml)
    except Exception as e:  # pragma: no cover - defensive
        logger.debug("theme-font resolution failed: %s", e)
    return {}


def _parse_theme_fonts(theme_xml: bytes) -> dict:
    """Parse the <a:majorFont>/<a:minorFont> Latin typefaces from theme1.xml."""
    from lxml import etree

    try:
        root = etree.fromstring(theme_xml)
    except Exception:
        return {}

    ns = {"a": _THEME_NS}
    out = {}
    major = root.find(".//a:fontScheme/a:majorFont/a:latin", ns)
    if major is not None and major.get("typeface"):
        out["major"] = major.get("typeface")
    minor = root.find(".//a:fontScheme/a:minorFont/a:latin", ns)
    if minor is not None and minor.get("typeface"):
        out["minor"] = minor.get("typeface")
    return out


def _extract_font_roles(doc, theme_fonts=None) -> dict:
    """Aggregate font usage grouped by paragraph style (role).

    E.g., text in paragraphs styled "Heading 1" vs "Normal" vs "Body Text"
    typically uses different fonts/sizes — we produce one summary per role
    rather than a single global histogram.
    """
    from docx.shared import Pt

    theme_fonts = theme_fonts or {}
    # When a run inherits from theme, python-docx shows name=None and we
    # fall back to the theme major/minor Latin fonts we resolved earlier.
    default_body = theme_fonts.get("minor")
    default_heading = theme_fonts.get("major") or default_body

    role_fonts: dict[str, Counter] = defaultdict(Counter)
    role_totals: Counter = Counter()

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        style_name = (para.style.name if para.style else "Normal") or "Normal"
        role_key = _canonicalize_role(style_name)
        role_totals[role_key] += 1
        # Role-based default for runs that don't specify a font name
        role_default = (
            default_heading if role_key.startswith(("Heading", "Title")) else default_body
        )
        for run in para.runs:
            f = run.font
            name = f.name or role_default or "Default"
            size_pt = None
            try:
                if f.size is not None:
                    size_pt = round(f.size / Pt(1), 1)
            except Exception:
                pass
            key = (name, size_pt, bool(f.bold), bool(f.italic))
            role_fonts[role_key][key] += len(run.text or "")

    out = {}
    for role, counter in role_fonts.items():
        if not counter:
            continue
        (name, size_pt, bold, italic), count = counter.most_common(1)[0]
        out[role] = {
            "name": name,
            "size_pt": size_pt,
            "bold": bold,
            "italic": italic,
            "sample_char_count": count,
            "paragraph_count": role_totals[role],
        }
    return out


def _canonicalize_role(style_name: str) -> str:
    """Map verbose style names down to friendly role keys."""
    name = style_name.strip()
    lower = name.lower()
    if lower in ("normal", "body text", "default paragraph font", ""):
        return "body"
    if lower.startswith("heading"):
        # "Heading 1", "Heading 2 Char" → "Heading 1", "Heading 2"
        m = re.search(r"heading\s*(\d+)", lower)
        if m:
            return f"Heading {m.group(1)}"
        return "Heading"
    if lower.startswith("title"):
        return "Title"
    if "caption" in lower:
        return "Caption"
    if "footer" in lower:
        return "Footer"
    if "header" in lower:
        return "Header"
    if "list" in lower:
        return "List"
    if "toc" in lower:
        return "TOC"
    return name  # preserve custom style names verbatim


# ── Paragraph rhythm (spacing, indent, line-height) per role ───────

def _extract_paragraph_rhythm(doc) -> dict:
    """For each paragraph role, find the dominant rhythm values.

    Returns one entry per role:
        {
          "body": {
            "space_before_pt": 0,
            "space_after_pt": 6,
            "line_spacing": 1.15,
            "left_indent_inches": 0,
            "first_line_indent_inches": 0,
            "alignment": "left",
            "sample_size": 142
          },
          "Heading 1": { ... },
          ...
        }

    Values are the most-common setting across paragraphs of that role so a
    downstream generator can reproduce the document's whitespace feel, not
    just its fonts.
    """
    from docx.shared import Pt, Emu

    role_rhythm: dict[str, Counter] = defaultdict(Counter)

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        style_name = (para.style.name if para.style else "Normal") or "Normal"
        role_key = _canonicalize_role(style_name)
        pf = para.paragraph_format

        def _pt(v):
            try:
                return round(v / Pt(1), 1) if v else 0
            except Exception:
                return None

        def _in(v):
            try:
                return round(v / Emu(914400), 3) if v else 0
            except Exception:
                return None

        line = pf.line_spacing
        if line is not None and line > 4:  # stored in twips when > 4 (EMU-ish)
            try:
                line = round(line / 12, 2)  # normalize "exact" spacing
            except Exception:
                line = None

        record = (
            _pt(pf.space_before),
            _pt(pf.space_after),
            line if line is None else round(float(line), 2),
            _in(pf.left_indent),
            _in(pf.first_line_indent),
            str(pf.alignment) if pf.alignment is not None else None,
        )
        role_rhythm[role_key][record] += 1

    out = {}
    for role, counter in role_rhythm.items():
        if not counter:
            continue
        (sb, sa, line, left, first_line, align), count = counter.most_common(1)[0]
        out[role] = {
            "space_before_pt": sb,
            "space_after_pt": sa,
            "line_spacing": line,
            "left_indent_inches": left,
            "first_line_indent_inches": first_line,
            "alignment": _clean_alignment(align),
            "sample_size": count,
        }
    return out


def _clean_alignment(a: Optional[str]) -> Optional[str]:
    """Turn python-docx's ``WD_ALIGN_PARAGRAPH.CENTER (1)`` into "center"."""
    if not a:
        return None
    s = str(a).split(".")[-1].split(" ")[0].lower()
    mapping = {
        "left": "left", "right": "right", "center": "center",
        "justify": "justify", "justify_low": "justify",
        "distribute": "distribute",
    }
    return mapping.get(s, s)


# ── Numbering definitions (numbering.xml) ──────────────────────────

def _extract_numbering_definitions(doc) -> dict:
    """Walk numbering.xml and return a simplified map of list styles.

    Shape::

        {
          "lists": [
            {
              "numId": 1,
              "levels": [
                {"ilvl": 0, "format": "decimal", "text": "%1.", "indent_dxa": 720},
                {"ilvl": 1, "format": "lowerLetter", "text": "%2)", "indent_dxa": 1440},
              ]
            }
          ]
        }

    Enables a generator to reproduce *exactly* the numbering scheme a
    document uses ("1.  1.1.  1.1.1." vs "A. 1. i." vs bullets), which
    simply inferring from rendered text cannot do reliably.
    """
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        return {"lists": []}
    if numbering_part is None:
        return {"lists": []}

    from lxml import etree
    try:
        root = etree.fromstring(numbering_part.blob)
    except Exception:
        return {"lists": []}

    abstract_map: dict[str, list[dict]] = {}
    for abstract in root.findall(_qn("abstractNum")):
        aid = _attr(abstract, "abstractNumId")
        levels = []
        for lvl in abstract.findall(_qn("lvl")):
            ilvl = _attr(lvl, "ilvl")
            num_fmt_el = lvl.find(_qn("numFmt"))
            lvl_text_el = lvl.find(_qn("lvlText"))
            p_pr = lvl.find(_qn("pPr"))
            indent_dxa = None
            if p_pr is not None:
                ind = p_pr.find(_qn("ind"))
                if ind is not None:
                    left = _attr(ind, "left") or _attr(ind, "start")
                    try:
                        indent_dxa = int(left) if left else None
                    except ValueError:
                        indent_dxa = None
            levels.append({
                "ilvl": int(ilvl) if ilvl and ilvl.isdigit() else ilvl,
                "format": _attr(num_fmt_el, "val") if num_fmt_el is not None else None,
                "text": _attr(lvl_text_el, "val") if lvl_text_el is not None else None,
                "indent_dxa": indent_dxa,
            })
        abstract_map[aid] = levels

    lists_out = []
    for num in root.findall(_qn("num")):
        num_id = _attr(num, "numId")
        abstract_ref = num.find(_qn("abstractNumId"))
        abstract_id = _attr(abstract_ref, "val") if abstract_ref is not None else None
        if abstract_id in abstract_map:
            lists_out.append({
                "numId": int(num_id) if num_id and num_id.isdigit() else num_id,
                "levels": abstract_map[abstract_id][:6],  # first 6 levels is plenty
            })

    return {"lists": lists_out[:20]}  # cap for payload size


# ── Tables ─────────────────────────────────────────────────────────

def _extract_table(table, idx: int) -> dict:
    """Return a rich spec for a single ``docx.Table`` object."""
    tbl_el = table._tbl
    col_widths = _extract_col_widths(tbl_el)
    tbl_properties = _extract_tbl_properties(tbl_el)

    cells: list[list[dict]] = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(_extract_cell(cell))
        cells.append(row_data)

    # Detect the typical "header row" — when row 0 has shading on all cells,
    # or when <w:tblHeader/> is set on its <w:trPr>.
    has_header_row = _detect_header_row(table)

    return {
        "index": idx,
        "rows": len(table.rows),
        "cols": len(col_widths) or (len(table.columns) if table.rows else 0),
        "col_widths_dxa": col_widths,
        "has_header_row": has_header_row,
        "properties": tbl_properties,
        "cells": cells,
    }


def _extract_tbl_properties(tbl_el) -> dict:
    """Pull <w:tblPr> — borders, alignment, width, layout, cell margins."""
    tbl_pr = tbl_el.find(_qn("tblPr"))
    if tbl_pr is None:
        return {}

    out: dict = {}

    # Table width (w:tblW) — can be auto / dxa / pct
    tbl_w = tbl_pr.find(_qn("tblW"))
    if tbl_w is not None:
        w_type = _attr(tbl_w, "type")
        w_val = _attr(tbl_w, "w")
        if w_type:
            out["width"] = {"type": w_type, "value": int(w_val) if (w_val or "").lstrip("-").isdigit() else w_val}

    # Alignment on page (left / center / right)
    jc = tbl_pr.find(_qn("jc"))
    if jc is not None:
        out["alignment"] = _attr(jc, "val")

    # Table layout (fixed vs autofit)
    layout = tbl_pr.find(_qn("tblLayout"))
    if layout is not None:
        out["layout"] = _attr(layout, "type")

    # Table-wide borders
    borders = tbl_pr.find(_qn("tblBorders"))
    if borders is not None:
        b_out = {}
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = borders.find(_qn(side))
            if b is not None:
                b_out[side] = {
                    "val": _attr(b, "val"),
                    "sz": _attr(b, "sz"),
                    "color": _normalize_color(_attr(b, "color")),
                }
        if b_out:
            out["borders"] = b_out

    # Default cell margins (w:tblCellMar)
    cell_mar = tbl_pr.find(_qn("tblCellMar"))
    if cell_mar is not None:
        m_out = {}
        for side in ("top", "start", "bottom", "end", "left", "right"):
            m = cell_mar.find(_qn(side))
            if m is not None:
                val = _attr(m, "w")
                try:
                    m_out[f"{side}_dxa"] = int(val) if val else None
                except ValueError:
                    m_out[f"{side}_dxa"] = None
        if m_out:
            out["cell_margins_dxa"] = m_out

    # Table style reference (e.g. "GridTable1Light")
    tbl_style = tbl_pr.find(_qn("tblStyle"))
    if tbl_style is not None:
        out["style"] = _attr(tbl_style, "val")

    return out


def _extract_col_widths(tbl_el) -> list[int]:
    widths: list[int] = []
    grid = tbl_el.find(_qn("tblGrid"))
    if grid is None:
        return widths
    for col in grid.findall(_qn("gridCol")):
        w = _attr(col, "w")
        try:
            widths.append(int(w))
        except (TypeError, ValueError):
            widths.append(0)
    return widths


def _detect_header_row(table) -> bool:
    if not table.rows:
        return False
    first_row = table.rows[0]
    tr_pr = first_row._tr.find(_qn("trPr"))
    if tr_pr is not None and tr_pr.find(_qn("tblHeader")) is not None:
        return True
    # Fallback heuristic: every cell in row 0 has the same shading.
    shades = [_extract_cell_shading(c._tc) for c in first_row.cells]
    if shades and all(shades) and len(set(shades)) == 1:
        return True
    return False


def _extract_cell(cell) -> dict:
    """Extract a single cell's text + formatting metadata."""
    tc = cell._tc
    tc_pr = tc.find(_qn("tcPr"))
    text = (cell.text or "").strip()

    return {
        "text": text[:200],
        "shading": _extract_cell_shading(tc),
        "grid_span": _extract_grid_span(tc_pr),
        "v_merge": _extract_v_merge(tc_pr),
        "borders": _extract_cell_borders(tc_pr),
        "alignment": _extract_cell_alignment(tc_pr, cell),
    }


def _extract_cell_shading(tc_el) -> Optional[str]:
    """Return the fill color (hex, no #) or None."""
    tc_pr = tc_el.find(_qn("tcPr"))
    if tc_pr is None:
        return None
    shd = tc_pr.find(_qn("shd"))
    if shd is None:
        return None
    fill = _attr(shd, "fill")
    if fill in (None, "auto", ""):
        return None
    return fill.upper()


def _normalize_color(val: Optional[str]) -> Optional[str]:
    """Normalize a color attribute — uppercase hex, but preserve the special
    ``auto`` keyword (which means "inherit" / default black in OOXML)."""
    if not val:
        return None
    low = val.lower()
    if low in ("auto", "none", "nil"):
        return low
    return val.upper()


def _extract_grid_span(tc_pr) -> int:
    if tc_pr is None:
        return 1
    gs = tc_pr.find(_qn("gridSpan"))
    if gs is None:
        return 1
    val = _attr(gs, "val")
    try:
        return int(val) if val else 1
    except ValueError:
        return 1


def _extract_v_merge(tc_pr) -> Optional[str]:
    """Return 'restart', 'continue', or None."""
    if tc_pr is None:
        return None
    vm = tc_pr.find(_qn("vMerge"))
    if vm is None:
        return None
    return _attr(vm, "val") or "continue"


def _extract_cell_borders(tc_pr) -> dict:
    if tc_pr is None:
        return {}
    borders = tc_pr.find(_qn("tcBorders"))
    if borders is None:
        return {}
    out = {}
    for side in ("top", "left", "bottom", "right"):
        b = borders.find(_qn(side))
        if b is not None:
            out[side] = {
                "val": _attr(b, "val"),
                "sz": _attr(b, "sz"),
                "color": _normalize_color(_attr(b, "color")),
            }
    return out


def _extract_cell_alignment(tc_pr, cell) -> dict:
    out = {}
    if tc_pr is not None:
        v_align = tc_pr.find(_qn("vAlign"))
        if v_align is not None:
            out["vertical"] = _attr(v_align, "val")
    # Horizontal alignment is per-paragraph; take the first paragraph's
    # <w:jc val="..."/> if present as a hint.
    try:
        first_para = cell.paragraphs[0]
        jc = first_para._p.find(_qn("pPr"))
        if jc is not None:
            jc_el = jc.find(_qn("jc"))
            if jc_el is not None:
                out["horizontal"] = _attr(jc_el, "val")
    except (IndexError, AttributeError):
        pass
    return out


# ── Shading palette across all tables ──────────────────────────────

def _extract_shading_palette(doc) -> list[dict]:
    """Histogram of unique cell-shading colors used across all tables.

    Useful signal — e.g., ``[{"color":"BFBFBF","count":12}]`` tells the
    analyzer "this org uses BFBFBF as its section-header grey".
    """
    counter: Counter = Counter()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                shd = _extract_cell_shading(cell._tc)
                if shd:
                    counter[shd] += 1
    return [{"color": c, "count": n} for c, n in counter.most_common(10)]
