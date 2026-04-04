"""Low-level OOXML manipulation helpers for Word document generation.

Provides functions to set cell shading, table widths, page orientation,
flowchart shapes, and other formatting that python-docx doesn't expose.
"""

from lxml import etree
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement


# Namespace map for OOXML elements
WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
DML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
WPG_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP14_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
V_NS = "urn:schemas-microsoft-com:vml"


def set_cell_shading(cell, fill_color: str, val: str = "clear"):
    """Apply background fill color to a table cell.

    Args:
        cell: python-docx TableCell object
        fill_color: Hex color string (e.g. 'BFBFBF')
        val: Shading pattern value (default 'clear' for solid fill)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shading
    for existing_shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing_shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), val)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def set_cell_vertical_align(cell, align: str = "center"):
    """Set vertical alignment of a table cell.

    Args:
        cell: python-docx TableCell object
        align: 'top', 'center', or 'bottom'
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(existing)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def set_cell_width(cell, width_dxa: int):
    """Set the width of a table cell in DXA units."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcW")):
        tcPr.remove(existing)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_cell_margins(cell, top: int = 0, right: int = 0, bottom: int = 0, left: int = 0):
    """Set cell internal margins/padding in DXA."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(existing)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("start", left), ("bottom", bottom), ("end", right)]:
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:w"), str(val))
        elem.set(qn("w:type"), "dxa")
        tcMar.append(elem)
    tcPr.append(tcMar)


def set_cell_gridspan(cell, span: int):
    """Set gridSpan on a cell for column merging."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:gridSpan")):
        tcPr.remove(existing)
    gridSpan = OxmlElement("w:gridSpan")
    gridSpan.set(qn("w:val"), str(span))
    tcPr.append(gridSpan)


def set_cell_vmerge(cell, restart: bool = False):
    """Set vertical merge on a cell.

    Args:
        cell: python-docx TableCell
        restart: True = start of merge, False = continue merge
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:vMerge")):
        tcPr.remove(existing)
    vMerge = OxmlElement("w:vMerge")
    if restart:
        vMerge.set(qn("w:val"), "restart")
    tcPr.append(vMerge)


def set_table_width(table, width_dxa: int):
    """Set the total table width in DXA units."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def set_table_fixed_layout(table):
    """Set table layout to fixed (prevents auto-resizing)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def set_table_borders(table, size: int = 4, color: str = "auto"):
    """Set all table borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tblPr.append(borders)


def set_table_indent(table, indent_dxa: int):
    """Set table left indent."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    for existing in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(existing)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(indent_dxa))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)


def set_column_widths(table, widths: list[int]):
    """Set grid column widths for a table.

    Args:
        table: python-docx Table object
        widths: List of column widths in DXA
    """
    tbl = table._tbl
    # Remove existing tblGrid
    for existing in tbl.findall(qn("w:tblGrid")):
        tbl.remove(existing)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        tblGrid.append(gridCol)
    # Insert after tblPr
    tblPr = tbl.tblPr
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)


def set_paragraph_spacing(paragraph, before: int = None, after: int = None,
                          line: int = None, line_rule: str = None):
    """Set paragraph spacing in twips.

    Args:
        paragraph: python-docx Paragraph object
        before: Space before in twips
        after: Space after in twips
        line: Line spacing in twips (240 = single)
        line_rule: 'auto', 'exact', 'atLeast'
    """
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:spacing")):
        pPr.remove(existing)
    spacing = OxmlElement("w:spacing")
    if before is not None:
        spacing.set(qn("w:before"), str(before))
    if after is not None:
        spacing.set(qn("w:after"), str(after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
    if line_rule is not None:
        spacing.set(qn("w:lineRule"), line_rule)
    pPr.append(spacing)


def set_row_height(row, height_dxa: int, rule: str = "exact"):
    """Set table row height.

    Args:
        row: python-docx TableRow
        height_dxa: Height in DXA
        rule: 'exact', 'atLeast', 'auto'
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for existing in trPr.findall(qn("w:trHeight")):
        trPr.remove(existing)
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(height_dxa))
    trHeight.set(qn("w:hRule"), rule)
    trPr.append(trHeight)


def set_row_cant_split(row):
    """Prevent a table row from splitting across pages."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)


def add_page_number_field(paragraph):
    """Insert a PAGE field code into a paragraph for current page number."""
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fldChar_end)


def add_num_pages_field(paragraph):
    """Insert a NUMPAGES field code for total page count."""
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " NUMPAGES "
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fldChar_end)


def add_checkbox(paragraph, checked: bool = False):
    """Add a checkbox character to a paragraph.

    Uses Unicode ballot box characters for maximum compatibility.
    """
    run = paragraph.add_run()
    run.text = "\u2611" if checked else "\u2610"  # Checked / Unchecked box
    return run


def set_page_landscape(section):
    """Set a document section to landscape US Letter orientation.

    In the OOXML spec, landscape swaps width and height in pgSz
    and sets the orient attribute.
    """
    sectPr = section._sectPr
    pgSz = sectPr.find(qn("w:pgSz"))
    if pgSz is None:
        pgSz = OxmlElement("w:pgSz")
        sectPr.append(pgSz)
    # Landscape: width > height
    pgSz.set(qn("w:w"), "15840")
    pgSz.set(qn("w:h"), "12240")
    pgSz.set(qn("w:orient"), "landscape")


def set_page_margins(section, top: int = 1080, right: int = 720,
                     bottom: int = 720, left: int = 720,
                     header: int = 1080, footer: int = 1440, gutter: int = 0):
    """Set page margins for a document section in DXA."""
    sectPr = section._sectPr
    pgMar = sectPr.find(qn("w:pgMar"))
    if pgMar is None:
        pgMar = OxmlElement("w:pgMar")
        sectPr.append(pgMar)
    pgMar.set(qn("w:top"), str(top))
    pgMar.set(qn("w:right"), str(right))
    pgMar.set(qn("w:bottom"), str(bottom))
    pgMar.set(qn("w:left"), str(left))
    pgMar.set(qn("w:header"), str(header))
    pgMar.set(qn("w:footer"), str(footer))
    pgMar.set(qn("w:gutter"), str(gutter))


def set_cell_border(cell, side: str, val: str = "single", sz: int = 4,
                    color: str = "auto", space: int = 0):
    """Set a border on a specific side of a cell.

    Args:
        cell: python-docx TableCell
        side: 'top', 'bottom', 'left', 'right'
        val: Border style ('single', 'double', 'none', etc.)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for existing in tcBorders.findall(qn(f"w:{side}")):
        tcBorders.remove(existing)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), val)
    border.set(qn("w:sz"), str(sz))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    tcBorders.append(border)


# ── Flowchart OOXML Drawing Shapes ──

EMU_PER_INCH = 914400
EMU_PER_CM = 360000
EMU_PER_PT = 12700


def _make_emu(inches: float = 0, cm: float = 0, pt: float = 0, emu: int = 0) -> int:
    """Convert measurement to EMU (English Metric Units)."""
    return int(inches * EMU_PER_INCH + cm * EMU_PER_CM + pt * EMU_PER_PT + emu)


def build_shape_rect(x: int, y: int, w: int, h: int, text: str,
                     fill_color: str = "4472C4", line_color: str = "2F5597",
                     font_size_pt: int = 9) -> etree._Element:
    """Build an OOXML rectangle shape (wps:wsp) for flowcharts.

    Args:
        x, y: Position in EMU
        w, h: Size in EMU
        text: Text content inside the shape
        fill_color: Hex fill color
        line_color: Hex outline color
        font_size_pt: Font size in points

    Returns:
        lxml Element representing wps:wsp
    """
    nsmap = {
        "wps": WPS_NS,
        "a": DML_NS,
    }

    wsp = etree.SubElement(etree.Element("tmp"), f"{{{WPS_NS}}}wsp")
    wsp = etree.Element(f"{{{WPS_NS}}}wsp", nsmap=nsmap)

    # Non-visual properties
    cNvSpPr = etree.SubElement(wsp, f"{{{WPS_NS}}}cNvSpPr")

    # Shape properties
    spPr = etree.SubElement(wsp, f"{{{WPS_NS}}}spPr")

    # Transform
    xfrm = etree.SubElement(spPr, f"{{{DML_NS}}}xfrm")
    off = etree.SubElement(xfrm, f"{{{DML_NS}}}off")
    off.set("x", str(x))
    off.set("y", str(y))
    ext = etree.SubElement(xfrm, f"{{{DML_NS}}}ext")
    ext.set("cx", str(w))
    ext.set("cy", str(h))

    # Preset geometry - rectangle
    prstGeom = etree.SubElement(spPr, f"{{{DML_NS}}}prstGeom")
    prstGeom.set("prst", "rect")
    etree.SubElement(prstGeom, f"{{{DML_NS}}}avLst")

    # Fill
    solidFill = etree.SubElement(spPr, f"{{{DML_NS}}}solidFill")
    srgbClr = etree.SubElement(solidFill, f"{{{DML_NS}}}srgbClr")
    srgbClr.set("val", fill_color)

    # Outline
    ln = etree.SubElement(spPr, f"{{{DML_NS}}}ln")
    ln.set("w", "12700")  # 1pt
    solidFill2 = etree.SubElement(ln, f"{{{DML_NS}}}solidFill")
    srgbClr2 = etree.SubElement(solidFill2, f"{{{DML_NS}}}srgbClr")
    srgbClr2.set("val", line_color)

    # Text body
    txBody = etree.SubElement(wsp, f"{{{WPS_NS}}}txBody")
    bodyPr = etree.SubElement(txBody, f"{{{DML_NS}}}bodyPr")
    bodyPr.set("vertOverflow", "clip")
    bodyPr.set("wrap", "square")
    bodyPr.set("anchor", "ctr")
    etree.SubElement(txBody, f"{{{DML_NS}}}lstStyle")

    p = etree.SubElement(txBody, f"{{{DML_NS}}}p")
    pPr = etree.SubElement(p, f"{{{DML_NS}}}pPr")
    pPr.set("algn", "ctr")
    r = etree.SubElement(p, f"{{{DML_NS}}}r")
    rPr = etree.SubElement(r, f"{{{DML_NS}}}rPr")
    rPr.set("lang", "en-US")
    rPr.set("sz", str(font_size_pt * 100))
    solidFill3 = etree.SubElement(rPr, f"{{{DML_NS}}}solidFill")
    srgbClr3 = etree.SubElement(solidFill3, f"{{{DML_NS}}}srgbClr")
    srgbClr3.set("val", "FFFFFF")
    t = etree.SubElement(r, f"{{{DML_NS}}}t")
    t.text = text

    return wsp


def build_shape_diamond(x: int, y: int, w: int, h: int, text: str,
                        fill_color: str = "ED7D31", line_color: str = "C55A11",
                        font_size_pt: int = 8) -> etree._Element:
    """Build an OOXML diamond shape for decision nodes."""
    wsp = build_shape_rect(x, y, w, h, text, fill_color, line_color, font_size_pt)
    # Change preset geometry to diamond
    spPr = wsp.find(f"{{{WPS_NS}}}spPr")
    prstGeom = spPr.find(f"{{{DML_NS}}}prstGeom")
    prstGeom.set("prst", "diamond")
    return wsp


def build_shape_oval(x: int, y: int, w: int, h: int, text: str,
                     fill_color: str = "70AD47", line_color: str = "548235",
                     font_size_pt: int = 9) -> etree._Element:
    """Build an OOXML oval/ellipse shape for start/end nodes."""
    wsp = build_shape_rect(x, y, w, h, text, fill_color, line_color, font_size_pt)
    spPr = wsp.find(f"{{{WPS_NS}}}spPr")
    prstGeom = spPr.find(f"{{{DML_NS}}}prstGeom")
    prstGeom.set("prst", "ellipse")
    return wsp


def build_connector_shape(x1: int, y1: int, x2: int, y2: int,
                          line_color: str = "404040") -> etree._Element:
    """Build a straight arrow connector between two points."""
    nsmap = {
        "wps": WPS_NS,
        "a": DML_NS,
    }
    wsp = etree.Element(f"{{{WPS_NS}}}wsp", nsmap=nsmap)
    cNvCnPr = etree.SubElement(wsp, f"{{{WPS_NS}}}cNvCnPr")

    spPr = etree.SubElement(wsp, f"{{{WPS_NS}}}spPr")

    # Transform
    xfrm = etree.SubElement(spPr, f"{{{DML_NS}}}xfrm")
    if y2 < y1:
        xfrm.set("flipV", "1")
    off = etree.SubElement(xfrm, f"{{{DML_NS}}}off")
    off.set("x", str(min(x1, x2)))
    off.set("y", str(min(y1, y2)))
    ext = etree.SubElement(xfrm, f"{{{DML_NS}}}ext")
    ext.set("cx", str(abs(x2 - x1) or 1))
    ext.set("cy", str(abs(y2 - y1) or 1))

    # Preset geometry - straight connector
    prstGeom = etree.SubElement(spPr, f"{{{DML_NS}}}prstGeom")
    prstGeom.set("prst", "straightConnector1")
    etree.SubElement(prstGeom, f"{{{DML_NS}}}avLst")

    # Line style with arrow
    ln = etree.SubElement(spPr, f"{{{DML_NS}}}ln")
    ln.set("w", "12700")
    solidFill = etree.SubElement(ln, f"{{{DML_NS}}}solidFill")
    srgbClr = etree.SubElement(solidFill, f"{{{DML_NS}}}srgbClr")
    srgbClr.set("val", line_color)
    tailEnd = etree.SubElement(ln, f"{{{DML_NS}}}tailEnd")
    tailEnd.set("type", "triangle")
    tailEnd.set("w", "med")
    tailEnd.set("len", "med")

    return wsp


def build_flowchart_drawing(shapes: list, total_width_emu: int,
                            total_height_emu: int) -> etree._Element:
    """Wrap flowchart shapes in a complete wp:inline drawing element.

    Args:
        shapes: List of wps:wsp elements
        total_width_emu: Total canvas width in EMU
        total_height_emu: Total canvas height in EMU

    Returns:
        lxml Element for w:drawing containing the flowchart group
    """
    drawing = OxmlElement("w:drawing")

    inline = etree.SubElement(drawing, f"{{{WP_NS}}}inline")
    inline.set("distT", "0")
    inline.set("distB", "0")
    inline.set("distL", "0")
    inline.set("distR", "0")

    extent = etree.SubElement(inline, f"{{{WP_NS}}}extent")
    extent.set("cx", str(total_width_emu))
    extent.set("cy", str(total_height_emu))

    docPr = etree.SubElement(inline, f"{{{WP_NS}}}docPr")
    docPr.set("id", "100")
    docPr.set("name", "Flowchart")

    graphic = etree.SubElement(inline, f"{{{DML_NS}}}graphic")
    graphicData = etree.SubElement(graphic, f"{{{DML_NS}}}graphicData")
    graphicData.set("uri", WPG_NS)

    # Group shape
    wgp = etree.SubElement(graphicData, f"{{{WPG_NS}}}wgp")
    cNvGrpSpPr = etree.SubElement(wgp, f"{{{WPG_NS}}}cNvGrpSpPr")

    grpSpPr = etree.SubElement(wgp, f"{{{WPG_NS}}}grpSpPr")
    xfrm = etree.SubElement(grpSpPr, f"{{{DML_NS}}}xfrm")
    off = etree.SubElement(xfrm, f"{{{DML_NS}}}off")
    off.set("x", "0")
    off.set("y", "0")
    ext = etree.SubElement(xfrm, f"{{{DML_NS}}}ext")
    ext.set("cx", str(total_width_emu))
    ext.set("cy", str(total_height_emu))
    chOff = etree.SubElement(xfrm, f"{{{DML_NS}}}chOff")
    chOff.set("x", "0")
    chOff.set("y", "0")
    chExt = etree.SubElement(xfrm, f"{{{DML_NS}}}chExt")
    chExt.set("cx", str(total_width_emu))
    chExt.set("cy", str(total_height_emu))

    # Add all shapes to the group
    for shape in shapes:
        wgp.append(shape)

    return drawing
