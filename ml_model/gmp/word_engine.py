"""Advanced DOCX generation engine for GMP documents.

Produces Word documents matching exact pharmaceutical/biotech formatting:
landscape orientation, complex header/footer tables, gray-shaded sections,
step procedures with 4-5 column layouts, checkboxes, and flowcharts.
"""

import io
import os
import logging
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Twips, Emu
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .template_schema import (
    DocumentTemplate, SectionDefinition, SectionType,
    TableConfig, StepProcedureConfig,
)
from . import ooxml_helpers as ox

logger = logging.getLogger(__name__)

def _looks_like_hex(val) -> bool:
    """Return True if ``val`` is a 6-char RGB hex string (no leading #)."""
    if not isinstance(val, str):
        return False
    v = val.lstrip("#").strip()
    return len(v) == 6 and all(c in "0123456789abcdefABCDEF" for c in v)


# Formatting defaults — used when no account_style override is supplied.
# Values match the Fred Hutch Cell Processing Facility example.
TABLE_WIDTH = 14580
HEADER_FILL = "BFBFBF"
LABEL_FILL = "F2F2F2"
DEFAULT_FONT = "Calibri"
BODY_SIZE = Pt(10)
SECTION_HEADER_SIZE = Pt(11)
FOOTER_SIZE = Pt(8)


class GMPWordEngine:
    """Generates GMP-compliant Word documents from template + data.

    Style customisation:
        Call ``generate(template, data, account_style=…)`` with a dict
        produced by ``style_consolidation.consolidate_style`` to override
        the defaults — fonts, shading colors, page size, margins — to
        match a specific customer's uploaded documents.

        When ``account_style`` is missing or a particular key is null,
        the Fred-Hutch defaults above apply.
    """

    def __init__(self):
        self.doc = None
        # These get re-bound by ``_apply_account_style`` on every call;
        # module constants are the fallback.
        self._default_font = DEFAULT_FONT
        self._body_size = BODY_SIZE
        self._header_fill = HEADER_FILL
        self._label_fill = LABEL_FILL
        self._account_page = None  # dict from consolidated style (may be None)

    def generate(
        self,
        template: DocumentTemplate,
        data: dict,
        account_style: Optional[dict] = None,
    ) -> bytes:
        """Generate a complete DOCX document.

        Args:
            template: Document template with formatting rules and sections
            data: Dict with section data to populate the template
            account_style: Optional consolidated style dict from
                ``style_consolidation.consolidate_style``. When present,
                overrides the default font, body size, header/label
                shading, and page setup.

        Returns:
            DOCX file as bytes
        """
        self.doc = Document()
        self._apply_account_style(account_style)
        self._setup_styles()
        self._setup_page(template)
        self._build_header(template, data)
        self._build_footer(template, data)

        # Build each section
        for section_def in template.sections:
            section_data = data.get(section_def.id, {})
            self._build_section(section_def, section_data, data)

        # Write to bytes
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _apply_account_style(self, account_style: Optional[dict]):
        """Set instance attributes from a consolidated account style.

        Falls back to module-level defaults for any key that's null or
        missing from ``account_style``.
        """
        self._default_font = DEFAULT_FONT
        self._body_size = BODY_SIZE
        self._header_fill = HEADER_FILL
        self._label_fill = LABEL_FILL
        self._account_page = None

        if not isinstance(account_style, dict):
            return

        body_font = account_style.get("body_font") or {}
        if body_font.get("name"):
            self._default_font = body_font["name"]
        if body_font.get("size_pt"):
            try:
                self._body_size = Pt(float(body_font["size_pt"]))
            except (TypeError, ValueError):
                pass

        shading = account_style.get("shading_roles") or {}
        if _looks_like_hex(shading.get("section_header")):
            self._header_fill = shading["section_header"].upper()
        if _looks_like_hex(shading.get("label_cell")):
            self._label_fill = shading["label_cell"].upper()

        page = account_style.get("page") or {}
        if page:
            self._account_page = page

    def _setup_styles(self):
        """Configure default document styles from the (possibly overridden)
        body-font attributes."""
        style = self.doc.styles["Normal"]
        style.font.name = self._default_font
        style.font.size = self._body_size
        style.paragraph_format.space_before = Twips(0)
        style.paragraph_format.space_after = Twips(0)

    def _setup_page(self, template: DocumentTemplate):
        """Set page orientation, size, and margins.

        When a consolidated ``_account_page`` is present, it wins over the
        template's defaults — the user's uploaded documents dictate layout.
        """
        section = self.doc.sections[0]

        page = self._account_page or {}
        orientation = page.get("orientation") or template.orientation.value
        size_inches = page.get("size_inches")
        margins_inches = page.get("margins_inches") or {}

        if orientation == "landscape":
            ox.set_page_landscape(section)

        # If the learned style declares explicit page dimensions, apply
        # them (in DXA — 1 inch = 1440 DXA).
        if size_inches and len(size_inches) == 2:
            sectPr = section._sectPr
            pgSz = sectPr.find(qn("w:pgSz"))
            if pgSz is None:
                pgSz = OxmlElement("w:pgSz")
                sectPr.append(pgSz)
            pgSz.set(qn("w:w"), str(int(size_inches[0] * 1440)))
            pgSz.set(qn("w:h"), str(int(size_inches[1] * 1440)))
        elif orientation != "landscape":
            sectPr = section._sectPr
            pgSz = sectPr.find(qn("w:pgSz"))
            if pgSz is None:
                pgSz = OxmlElement("w:pgSz")
                sectPr.append(pgSz)
            pgSz.set(qn("w:w"), str(template.page_size.width_dxa))
            pgSz.set(qn("w:h"), str(template.page_size.height_dxa))

        # Margins — prefer learned values when set, fall back per-side.
        m = template.margins

        def _dxa(inches):
            try:
                return int(float(inches) * 1440)
            except (TypeError, ValueError):
                return None

        ox.set_page_margins(
            section,
            top=_dxa(margins_inches.get("top")) or m.top,
            right=_dxa(margins_inches.get("right")) or m.right,
            bottom=_dxa(margins_inches.get("bottom")) or m.bottom,
            left=_dxa(margins_inches.get("left")) or m.left,
            header=m.header, footer=m.footer, gutter=m.gutter,
        )

    def _build_header(self, template: DocumentTemplate, data: dict):
        """Build the complex header table with logo, title, page numbers."""
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False

        hdr_cfg = template.header_config
        col_widths = hdr_cfg.column_widths
        num_cols = hdr_cfg.columns

        # Create header table
        table = header.add_table(rows=3, cols=num_cols, width=Twips(hdr_cfg.total_width))
        ox.set_table_width(table, hdr_cfg.total_width)
        ox.set_table_fixed_layout(table)
        ox.set_table_borders(table, size=4, color="auto")
        ox.set_column_widths(table, col_widths)

        doc_number = data.get("doc_number", "BR-XXX-XX")
        doc_title = data.get("doc_title", template.name)
        effective_date = data.get("effective_date", datetime.now().strftime("%d%b%Y").upper())
        revision = data.get("revision", "01")

        # Row 0: Logo | Title (center, bold 12pt) | Page X of Y | Doc # | Eff Date | Rev
        # Logo cell (spans 2 cols, 3 rows)
        logo_cell = table.cell(0, 0)
        ox.set_cell_gridspan(logo_cell, 2)
        ox.set_cell_vmerge(logo_cell, restart=True)
        ox.set_cell_vertical_align(logo_cell, "center")
        logo_p = logo_cell.paragraphs[0]
        logo_run = logo_p.add_run("[LOGO]")
        logo_run.font.size = Pt(8)
        logo_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Title cell (spans 1 col, 3 rows)
        title_cell = table.cell(0, 2)
        ox.set_cell_vmerge(title_cell, restart=True)
        ox.set_cell_vertical_align(title_cell, "center")
        title_p = title_cell.paragraphs[0]
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(doc_title)
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_run.font.name = self._default_font

        # Page number cell
        page_cell = table.cell(0, 3)
        ox.set_cell_vertical_align(page_cell, "center")
        page_p = page_cell.paragraphs[0]
        page_run = page_p.add_run("Page ")
        page_run.font.size = self._body_size
        page_run.font.name = self._default_font
        ox.add_page_number_field(page_p)
        of_run = page_p.add_run(" of ")
        of_run.font.size = self._body_size
        ox.add_num_pages_field(page_p)

        # Doc number label
        dn_label_cell = table.cell(0, 4)
        ox.set_cell_vertical_align(dn_label_cell, "center")
        dn_label_p = dn_label_cell.paragraphs[0]
        dn_label_run = dn_label_p.add_run("Document No:")
        dn_label_run.font.size = self._body_size
        dn_label_run.font.name = self._default_font

        # Doc number value
        dn_val_cell = table.cell(0, 5)
        ox.set_cell_vertical_align(dn_val_cell, "center")
        dn_val_p = dn_val_cell.paragraphs[0]
        dn_val_run = dn_val_p.add_run(doc_number)
        dn_val_run.font.size = self._body_size
        dn_val_run.font.name = self._default_font

        # Row 1: (logo cont) | (title cont) | Production ID | Eff Date label | Eff Date val
        for r in [1, 2]:
            # Continue vertical merge for logo and title
            ox.set_cell_vmerge(table.cell(r, 0))
            ox.set_cell_gridspan(table.cell(r, 0), 2)
            ox.set_cell_vmerge(table.cell(r, 2))

        # Production ID cell
        pid_cell = table.cell(1, 3)
        ox.set_cell_vertical_align(pid_cell, "center")
        pid_p = pid_cell.paragraphs[0]
        pid_run = pid_p.add_run("Production ID:")
        pid_run.font.size = self._body_size

        # Effective Date label
        ed_label_cell = table.cell(1, 4)
        ox.set_cell_vertical_align(ed_label_cell, "center")
        ed_p = ed_label_cell.paragraphs[0]
        ed_run = ed_p.add_run("Effective Date:")
        ed_run.font.size = self._body_size

        # Effective Date value
        ed_val_cell = table.cell(1, 5)
        ox.set_cell_vertical_align(ed_val_cell, "center")
        ed_val_p = ed_val_cell.paragraphs[0]
        ed_val_run = ed_val_p.add_run(effective_date)
        ed_val_run.font.size = self._body_size

        # Row 2: Lot Number | Revision label | Revision value
        lot_cell = table.cell(2, 3)
        ox.set_cell_vertical_align(lot_cell, "center")
        lot_p = lot_cell.paragraphs[0]
        lot_run = lot_p.add_run("Lot Number:")
        lot_run.font.size = self._body_size

        rev_label_cell = table.cell(2, 4)
        ox.set_cell_vertical_align(rev_label_cell, "center")
        rev_p = rev_label_cell.paragraphs[0]
        rev_run = rev_p.add_run("Revision No:")
        rev_run.font.size = self._body_size

        rev_val_cell = table.cell(2, 5)
        ox.set_cell_vertical_align(rev_val_cell, "center")
        rev_val_p = rev_val_cell.paragraphs[0]
        rev_val_run = rev_val_p.add_run(revision)
        rev_val_run.font.size = self._body_size

        # Set row heights
        for row in table.rows:
            ox.set_row_height(row, hdr_cfg.row_height_dxa, rule="exact")
            ox.set_row_cant_split(row)

    def _build_footer(self, template: DocumentTemplate, data: dict):
        """Build the footer with confidentiality statement."""
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        ftr_cfg = template.footer_config
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add top border
        if ftr_cfg.border_top:
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            top_bdr = OxmlElement("w:top")
            top_bdr.set(qn("w:val"), "single")
            top_bdr.set(qn("w:sz"), "4")
            top_bdr.set(qn("w:space"), "1")
            top_bdr.set(qn("w:color"), "auto")
            pBdr.append(top_bdr)
            pPr.append(pBdr)

        ox.set_paragraph_spacing(p, before=40)

        text = ftr_cfg.text
        if ftr_cfg.organization:
            text = f"{text} \u2013 {ftr_cfg.organization}"

        run = p.add_run(text)
        run.bold = ftr_cfg.bold
        run.font.size = Pt(ftr_cfg.font_size_half_pt / 2)
        run.font.name = self._default_font

    def _build_section(self, section_def: SectionDefinition, section_data: dict,
                       all_data: dict):
        """Dispatch to the appropriate section builder based on type."""
        builders = {
            SectionType.APPROVAL_BLOCK: self._build_approval_block,
            SectionType.TABLE: self._build_generic_table,
            SectionType.STEP_PROCEDURE: self._build_step_procedure,
            SectionType.REFERENCES: self._build_references_table,
            SectionType.ATTACHMENTS: self._build_attachments_table,
            SectionType.CHECKLIST: self._build_checklist,
            SectionType.FREE_TEXT: self._build_free_text,
            SectionType.EQUIPMENT_LIST: self._build_equipment_list,
            SectionType.MATERIALS_LIST: self._build_materials_list,
            SectionType.GENERAL_INSTRUCTIONS: self._build_general_instructions,
            SectionType.LABEL_ACCOUNTABILITY: self._build_label_accountability,
            SectionType.COMMENTS: self._build_comments,
            SectionType.REVIEW: self._build_review,
            SectionType.FLOWCHART: self._build_flowchart_section,
        }

        builder = builders.get(section_def.type)
        if builder:
            builder(section_def, section_data, all_data)
        else:
            logger.warning(f"No builder for section type: {section_def.type}")

    def _add_section_header(self, title: str, col_count: int = 8,
                            col_widths: Optional[list[int]] = None):
        """Add a gray-shaded section header row spanning full table width.

        Returns the table for further row additions.
        """
        widths = col_widths or self._default_col_widths(col_count)
        table = self.doc.add_table(rows=1, cols=col_count)
        ox.set_table_width(table, TABLE_WIDTH)
        ox.set_table_fixed_layout(table)
        ox.set_table_borders(table)
        ox.set_column_widths(table, widths)
        ox.set_table_indent(table, -5)

        # Header row
        header_cell = table.cell(0, 0)
        # Merge all columns
        for i in range(1, col_count):
            header_cell = header_cell.merge(table.cell(0, i))
        ox.set_cell_shading(header_cell, self._header_fill)
        ox.set_cell_vertical_align(header_cell, "center")

        p = header_cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)
        run = p.add_run(f"  {title}")
        run.bold = True
        run.font.size = SECTION_HEADER_SIZE
        run.font.name = self._default_font

        return table

    def _default_col_widths(self, col_count: int) -> list[int]:
        """Generate default column widths for a given column count."""
        if col_count == 8:
            return [2160, 3240, 720, 1260, 2070, 3060, 720, 1350]
        elif col_count == 5:
            return [5400, 2160, 3240, 1890, 1890]
        elif col_count == 4:
            return [5400, 2700, 3240, 3240]
        elif col_count == 2:
            return [7290, 7290]
        else:
            w = TABLE_WIDTH // col_count
            return [w] * (col_count - 1) + [TABLE_WIDTH - w * (col_count - 1)]

    def _add_label_value_row(self, table, label: str, value: str = "",
                             label_cols: int = 1, value_cols: int = 1,
                             label_fill: Optional[str] = None):
        """Add a label-value row to a table.

        ``label_fill`` defaults to the instance's current label shading
        (which may have been overridden by an account style).
        """
        if label_fill is None:
            label_fill = self._label_fill
        row = table.add_row()

        label_cell = row.cells[0]
        if label_cols > 1:
            for i in range(1, label_cols):
                label_cell = label_cell.merge(row.cells[i])
        ox.set_cell_shading(label_cell, label_fill)
        ox.set_cell_vertical_align(label_cell, "center")
        p = label_cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)
        run = p.add_run(label)
        run.font.size = self._body_size
        run.font.name = self._default_font

        value_cell = row.cells[label_cols]
        if value_cols > 1:
            for i in range(label_cols + 1, label_cols + value_cols):
                value_cell = value_cell.merge(row.cells[i])
        ox.set_cell_vertical_align(value_cell, "center")
        p = value_cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)
        if value:
            run = p.add_run(value)
            run.font.size = self._body_size
            run.font.name = self._default_font

        return row

    # ── Section Builders ──

    def _build_approval_block(self, section_def: SectionDefinition,
                              section_data: dict, all_data: dict):
        """Build the approval block with Document Owner, Technical Authority, etc."""
        table = self._add_section_header("APPROVED BY:")

        approvers = section_data.get("approvers", [
            {"role": "Document Owner", "name": "", "date": ""},
            {"role": "Technical Authority", "name": "", "date": ""},
            {"role": "TPP Facility Director", "name": "", "date": ""},
            {"role": "Quality Assurance", "name": "", "date": ""},
        ])

        for approver in approvers:
            row = table.add_row()
            # Role label (col 0)
            role_cell = row.cells[0]
            ox.set_cell_shading(role_cell, self._label_fill)
            ox.set_cell_vertical_align(role_cell, "center")
            p = role_cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(approver["role"] + ":")
            run.font.size = self._body_size
            run.font.name = self._default_font

            # Name (cols 1-2)
            name_cell = row.cells[1]
            name_cell = name_cell.merge(row.cells[2])
            ox.set_cell_vertical_align(name_cell, "center")
            p = name_cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            if approver.get("name"):
                run = p.add_run(approver["name"])
                run.font.size = self._body_size
                run.font.name = self._default_font

            # Date label (col 3)
            date_label_cell = row.cells[3]
            ox.set_cell_shading(date_label_cell, self._label_fill)
            ox.set_cell_vertical_align(date_label_cell, "center")
            p = date_label_cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run("Date:")
            run.font.size = self._body_size
            run.font.name = self._default_font

            # Date value (col 4)
            date_cell = row.cells[4]
            ox.set_cell_vertical_align(date_cell, "center")
            p = date_cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            if approver.get("date"):
                run = p.add_run(approver["date"])
                run.font.size = self._body_size
                run.font.name = self._default_font

        # Batch Record Issuance section
        issue_row = table.add_row()
        issue_cell = issue_row.cells[0]
        for i in range(1, 8):
            issue_cell = issue_cell.merge(issue_row.cells[i])
        ox.set_cell_shading(issue_cell, self._header_fill)
        p = issue_cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)
        run = p.add_run("  Batch Record issued for processing:")
        run.bold = True
        run.font.size = SECTION_HEADER_SIZE
        run.font.name = self._default_font

        # Issuance details
        issuance_items = [
            "Copy is an accurate reproduction of the current approved version.",
            "Production ID and Lot Number are recorded on all pages including attachments.",
        ]
        for item in issuance_items:
            row = table.add_row()
            cell = row.cells[0]
            for i in range(1, 8):
                cell = cell.merge(row.cells[i])
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(item)
            run.font.size = self._body_size
            run.font.name = self._default_font

        # Issued By Signature row
        row = table.add_row()
        cell = row.cells[0]
        for i in range(1, 8):
            cell = cell.merge(row.cells[i])
        p = cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)
        run = p.add_run("Issued By Signature / Date:")
        run.font.size = self._body_size
        run.font.name = self._default_font

    def _build_references_table(self, section_def: SectionDefinition,
                                section_data: dict, all_data: dict):
        """Build the references table with doc numbers and titles."""
        col_widths = [3240, 11340]
        table = self._add_section_header("REFERENCES", col_count=2,
                                         col_widths=col_widths)

        # Column headers
        row = table.add_row()
        for i, (header, width) in enumerate(zip(
            ["Document Number", "Title"], col_widths
        )):
            cell = row.cells[i]
            ox.set_cell_shading(cell, self._label_fill)
            ox.set_cell_vertical_align(cell, "center")
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(header)
            run.bold = True
            run.font.size = self._body_size
            run.font.name = self._default_font

        # Reference rows
        references = section_data.get("references", [])
        for ref in references:
            row = table.add_row()
            doc_cell = row.cells[0]
            ox.set_cell_vertical_align(doc_cell, "center")
            p = doc_cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(ref.get("doc_number", ""))
            run.font.size = self._body_size
            run.font.name = self._default_font

            title_cell = row.cells[1]
            ox.set_cell_vertical_align(title_cell, "center")
            p = title_cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(ref.get("title", ""))
            run.font.size = self._body_size
            run.font.name = self._default_font

    def _build_attachments_table(self, section_def: SectionDefinition,
                                 section_data: dict, all_data: dict):
        """Build the attachments table with doc numbers, titles, and quantities."""
        col_widths = [3240, 9900, 1440]
        table = self._add_section_header("ATTACHMENTS", col_count=3,
                                         col_widths=col_widths)

        # Column headers
        row = table.add_row()
        headers = ["Document Number", "Title", "Qty."]
        for i, header in enumerate(headers):
            cell = row.cells[i]
            ox.set_cell_shading(cell, self._label_fill)
            ox.set_cell_vertical_align(cell, "center")
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(header)
            run.bold = True
            run.font.size = self._body_size
            run.font.name = self._default_font

        attachments = section_data.get("attachments", [])
        for att in attachments:
            row = table.add_row()
            for i, key in enumerate(["doc_number", "title", "quantity"]):
                cell = row.cells[i]
                ox.set_cell_vertical_align(cell, "center")
                p = cell.paragraphs[0]
                ox.set_paragraph_spacing(p, before=60, after=60)
                val = str(att.get(key, ""))
                run = p.add_run(val)
                run.font.size = self._body_size
                run.font.name = self._default_font

    def _build_general_instructions(self, section_def: SectionDefinition,
                                    section_data: dict, all_data: dict):
        """Build the general instructions section."""
        table = self._add_section_header("GENERAL INSTRUCTIONS", col_count=1,
                                         col_widths=[TABLE_WIDTH])

        instructions = section_data.get("instructions", [])
        for instruction in instructions:
            row = table.add_row()
            cell = row.cells[0]
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(instruction)
            run.font.size = self._body_size
            run.font.name = self._default_font

    def _build_equipment_list(self, section_def: SectionDefinition,
                              section_data: dict, all_data: dict):
        """Build equipment list table."""
        col_widths = [7290, 7290]
        table = self._add_section_header("Equipment / materials list",
                                         col_count=2, col_widths=col_widths)

        # Equipment sub-header
        row = table.add_row()
        cell = row.cells[0]
        cell = cell.merge(row.cells[1])
        ox.set_cell_shading(cell, self._label_fill)
        p = cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)
        run = p.add_run("Equipment List:")
        run.bold = True
        run.font.size = self._body_size
        run.font.name = self._default_font

        # Equipment description header
        row = table.add_row()
        cell = row.cells[0]
        cell = cell.merge(row.cells[1])
        ox.set_cell_shading(cell, self._label_fill)
        p = cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)
        run = p.add_run("Description")
        run.bold = True
        run.font.size = self._body_size
        run.font.name = self._default_font

        equipment = section_data.get("equipment", [])
        for equip in equipment:
            row = table.add_row()
            cell = row.cells[0]
            cell = cell.merge(row.cells[1])
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            desc = equip if isinstance(equip, str) else equip.get("description", "")
            run = p.add_run(desc)
            run.font.size = self._body_size
            run.font.name = self._default_font

    def _build_materials_list(self, section_def: SectionDefinition,
                              section_data: dict, all_data: dict):
        """Build materials list table (side-by-side format)."""
        col_widths = [1800, 3600, 1890, 1800, 3600, 1890]
        table = self._add_section_header("Material List:", col_count=6,
                                         col_widths=col_widths)

        # Column headers (repeated for side-by-side)
        row = table.add_row()
        headers = ["CPF Part No.", "Material Description", "Quantity Required"] * 2
        for i, header in enumerate(headers):
            cell = row.cells[i]
            ox.set_cell_shading(cell, self._label_fill)
            ox.set_cell_vertical_align(cell, "center")
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = self._default_font

        materials = section_data.get("materials", [])
        # Split into two columns
        mid = (len(materials) + 1) // 2
        left_materials = materials[:mid]
        right_materials = materials[mid:]

        for i in range(max(len(left_materials), len(right_materials))):
            row = table.add_row()
            # Left side
            if i < len(left_materials):
                mat = left_materials[i]
                for j, key in enumerate(["part_number", "description", "quantity"]):
                    cell = row.cells[j]
                    ox.set_cell_vertical_align(cell, "center")
                    p = cell.paragraphs[0]
                    ox.set_paragraph_spacing(p, before=60, after=60)
                    run = p.add_run(str(mat.get(key, "")))
                    run.font.size = Pt(9)
                    run.font.name = self._default_font

            # Right side
            if i < len(right_materials):
                mat = right_materials[i]
                for j, key in enumerate(["part_number", "description", "quantity"]):
                    cell = row.cells[j + 3]
                    ox.set_cell_vertical_align(cell, "center")
                    p = cell.paragraphs[0]
                    ox.set_paragraph_spacing(p, before=60, after=60)
                    run = p.add_run(str(mat.get(key, "")))
                    run.font.size = Pt(9)
                    run.font.name = self._default_font

    def _build_step_procedure(self, section_def: SectionDefinition,
                              section_data: dict, all_data: dict):
        """Build day-by-day step procedure tables.

        Each processing day gets its own table with:
        - Section header (gray)
        - Column headers: Instruction | Variable | Result/Actual | Performed By | Checked By
        - Step rows with sub-steps, variables, calculations, checkboxes
        """
        step_cfg = section_def.step_config or StepProcedureConfig()
        col_widths = [col.width_dxa for col in step_cfg.columns]
        col_count = len(step_cfg.columns)

        title = section_data.get("title", section_def.title)
        table = self._add_section_header(title, col_count=col_count,
                                         col_widths=col_widths)

        # Column headers
        row = table.add_row()
        for i, col_def in enumerate(step_cfg.columns):
            cell = row.cells[i]
            ox.set_cell_shading(cell, step_cfg.label_fill)
            ox.set_cell_vertical_align(cell, "center")
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(col_def.title)
            run.bold = True
            run.font.size = self._body_size
            run.font.name = self._default_font

        # Step rows
        steps = section_data.get("steps", [])
        for step in steps:
            self._add_step_row(table, step, col_count)

        # Section review
        review_row = table.add_row()
        cell = review_row.cells[0]
        for i in range(1, col_count):
            cell = cell.merge(review_row.cells[i])
        ox.set_cell_shading(cell, self._label_fill)
        p = cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)
        run = p.add_run("Section Review")
        run.bold = True
        run.font.size = self._body_size
        run.font.name = self._default_font

        # MFG and QA review rows
        for reviewer in ["MFG Review (Initials/Date):", "QA Review (Initials/Date):"]:
            row = table.add_row()
            cell = row.cells[0]
            for i in range(1, col_count):
                cell = cell.merge(row.cells[i])
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(reviewer)
            run.font.size = self._body_size
            run.font.name = self._default_font

    def _add_step_row(self, table, step: dict, col_count: int):
        """Add a single step row to a procedure table."""
        row = table.add_row()

        # Instruction column
        instr_cell = row.cells[0]
        ox.set_cell_vertical_align(instr_cell, "top")
        p = instr_cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)

        # Step number + title
        step_num = step.get("number", "")
        step_title = step.get("title", "")
        if step_num:
            title_run = p.add_run(f"{step_num}  ")
            title_run.bold = True
            title_run.font.size = self._body_size
            title_run.font.name = self._default_font

        if step_title:
            title_run = p.add_run(step_title)
            title_run.bold = True
            title_run.font.size = self._body_size
            title_run.font.name = self._default_font

        # Sub-instructions
        instructions = step.get("instructions", [])
        for instr in instructions:
            p = instr_cell.add_paragraph()
            ox.set_paragraph_spacing(p, before=40, after=40)

            text = instr if isinstance(instr, str) else instr.get("text", "")
            is_bsc = False
            if isinstance(instr, dict):
                is_bsc = instr.get("bsc", False)

            if is_bsc:
                bsc_run = p.add_run("[BSC] ")
                bsc_run.bold = True
                bsc_run.font.size = self._body_size
                bsc_run.font.name = self._default_font

            run = p.add_run(text)
            run.font.size = self._body_size
            run.font.name = self._default_font

            # Handle verification options
            if isinstance(instr, dict) and instr.get("type") == "verification":
                for option in instr.get("options", []):
                    opt_p = instr_cell.add_paragraph()
                    ox.set_paragraph_spacing(opt_p, before=20, after=20)
                    ox.add_checkbox(opt_p, checked=False)
                    run = opt_p.add_run(f" {option}")
                    run.font.size = self._body_size
                    run.font.name = self._default_font

        # Variable column
        if col_count > 1:
            var_cell = row.cells[1]
            ox.set_cell_vertical_align(var_cell, "top")
            variables = step.get("variables", [])
            for var in variables:
                p = var_cell.add_paragraph() if var_cell.paragraphs[0].text else var_cell.paragraphs[0]
                ox.set_paragraph_spacing(p, before=40, after=40)
                var_name = var if isinstance(var, str) else var.get("name", "")
                run = p.add_run(var_name)
                run.font.size = self._body_size
                run.font.name = self._default_font

        # Result column (empty for filling in)
        if col_count > 2:
            result_cell = row.cells[2]
            ox.set_cell_vertical_align(result_cell, "top")
            results = step.get("results", [])
            for res in results:
                p = result_cell.add_paragraph() if result_cell.paragraphs[0].text else result_cell.paragraphs[0]
                ox.set_paragraph_spacing(p, before=40, after=40)
                run = p.add_run(res.get("unit", "") if isinstance(res, dict) else str(res))
                run.font.size = self._body_size
                run.font.name = self._default_font

    def _build_checklist(self, section_def: SectionDefinition,
                         section_data: dict, all_data: dict):
        """Build a review checklist with checkboxes."""
        table = self._add_section_header(section_def.title, col_count=1,
                                         col_widths=[TABLE_WIDTH])

        items = section_data.get("checklist_items", [])
        for item in items:
            row = table.add_row()
            cell = row.cells[0]
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            ox.add_checkbox(p, checked=False)
            text = item if isinstance(item, str) else item.get("text", "")
            run = p.add_run(f"  {text}")
            run.font.size = self._body_size
            run.font.name = self._default_font

            # Add N/A option if applicable
            if isinstance(item, dict) and item.get("na_option"):
                na_run = p.add_run(" (")
                na_run.font.size = self._body_size
                ox.add_checkbox(p, checked=False)
                na_run2 = p.add_run(" N/A)")
                na_run2.font.size = self._body_size

    def _build_comments(self, section_def: SectionDefinition,
                        section_data: dict, all_data: dict):
        """Build the comments section table."""
        col_widths = [2160, 9180, 3240]
        table = self._add_section_header("Comments", col_count=3,
                                         col_widths=col_widths)

        # Column headers
        row = table.add_row()
        for i, header in enumerate(["Step Reference", "Comments",
                                     "Recorded By (Initials/Date)"]):
            cell = row.cells[i]
            ox.set_cell_shading(cell, self._label_fill)
            ox.set_cell_vertical_align(cell, "center")
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(header)
            run.bold = True
            run.font.size = self._body_size
            run.font.name = self._default_font

        # Empty rows with N/A checkboxes
        num_rows = section_data.get("num_rows", 4)
        for _ in range(num_rows):
            row = table.add_row()
            ox.set_row_height(row, 504, rule="atLeast")
            for i in range(3):
                cell = row.cells[i]
                ox.set_cell_vertical_align(cell, "top")
                if i == 1:  # Comments column - add N/A checkbox
                    p = cell.paragraphs[0]
                    ox.set_paragraph_spacing(p, before=60, after=60)
                    ox.add_checkbox(p, checked=False)
                    run = p.add_run(" N/A")
                    run.font.size = self._body_size
                    run.font.name = self._default_font

    def _build_review(self, section_def: SectionDefinition,
                      section_data: dict, all_data: dict):
        """Build the manufacturing/QA review section."""
        review_type = section_data.get("review_type", "MANUFACTURING REVIEW")
        table = self._add_section_header(review_type, col_count=1,
                                         col_widths=[TABLE_WIDTH])

        # Checklist sub-header
        if section_data.get("checklist_items"):
            row = table.add_row()
            cell = row.cells[0]
            ox.set_cell_shading(cell, self._label_fill)
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run("Review Checklist")
            run.bold = True
            run.font.size = self._body_size
            run.font.name = self._default_font

            for item in section_data["checklist_items"]:
                row = table.add_row()
                cell = row.cells[0]
                p = cell.paragraphs[0]
                ox.set_paragraph_spacing(p, before=40, after=40)
                ox.add_checkbox(p, checked=False)
                text = item if isinstance(item, str) else item.get("text", "")
                run = p.add_run(f"  {text}")
                run.font.size = self._body_size
                run.font.name = self._default_font

        # Reviewer comments
        row = table.add_row()
        cell = row.cells[0]
        p = cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)
        run = p.add_run(f"{review_type.title()} Reviewer Comments:  ")
        run.font.size = self._body_size
        run.font.name = self._default_font
        ox.add_checkbox(p, checked=False)
        run2 = p.add_run("  N/A")
        run2.font.size = self._body_size
        run2.font.name = self._default_font

        # Signature line
        row = table.add_row()
        cell = row.cells[0]
        ox.set_row_height(row, 504, rule="atLeast")
        p = cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)
        run = p.add_run(f"{review_type.title()} Review Signature/Date: ")
        run.font.size = self._body_size
        run.font.name = self._default_font

    def _build_label_accountability(self, section_def: SectionDefinition,
                                    section_data: dict, all_data: dict):
        """Build the label accountability section."""
        col_widths = [5400, 2160, 3240, 1890, 1890]
        table = self._add_section_header("Label Accountability",
                                         col_count=5, col_widths=col_widths)

        # Column headers
        row = table.add_row()
        headers = ["Instruction", "Variable", "Result / Actual",
                   "Performed By (Initials/Date)", "Verified By (Initials/Date)"]
        for i, header in enumerate(headers):
            cell = row.cells[i]
            ox.set_cell_shading(cell, self._label_fill)
            ox.set_cell_vertical_align(cell, "center")
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = self._default_font

        # In-process labels row
        row = table.add_row()
        cell = row.cells[0]
        p = cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)
        run = p.add_run("In-Process Labels")
        run.bold = True
        run.font.size = self._body_size
        run.font.name = self._default_font
        p2 = cell.add_paragraph()
        ox.set_paragraph_spacing(p2, before=40, after=40)
        run2 = p2.add_run("Deface and discard any unused in-process labels.")
        run2.font.size = self._body_size
        run2.font.name = self._default_font

        # Infusion labels row
        row = table.add_row()
        cell = row.cells[0]
        p = cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)
        run = p.add_run("Infusion Labels")
        run.bold = True
        run.font.size = self._body_size
        run.font.name = self._default_font
        p2 = cell.add_paragraph()
        run2 = p2.add_run(
            "Do not discard remaining infusion labels. "
            "Deface unused Infusion Labels. "
            "Attach unused Infusion Labels to this record."
        )
        run2.font.size = self._body_size
        run2.font.name = self._default_font

        # Label count rows
        label_rows = [
            "Attach unused Infusion Labels to this record. Record Number attached.",
            "Record Number of Labels Printed",
            "Record Number of Labels Adhered to Record",
            "Record Number of Labels Used During Processing",
        ]
        for text in label_rows:
            row = table.add_row()
            cell = row.cells[0]
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(text)
            run.font.size = self._body_size
            run.font.name = self._default_font

            # Variable column - "Labels"
            var_cell = row.cells[2]
            p = var_cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run("Labels")
            run.font.size = self._body_size
            run.font.name = self._default_font

        # Accountability calculation row
        row = table.add_row()
        cell = row.cells[0]
        p = cell.paragraphs[0]
        ox.set_paragraph_spacing(p, before=60, after=60)
        run = p.add_run("Calculate: Label Accountability = Printed - Adhered - Used - Attached")
        run.font.size = self._body_size
        run.font.name = self._default_font

    def _build_generic_table(self, section_def: SectionDefinition,
                             section_data: dict, all_data: dict):
        """Build a generic table section from column definitions."""
        columns = section_def.columns
        if not columns:
            return

        col_widths = [col.width_dxa for col in columns]
        table = self._add_section_header(section_def.title,
                                         col_count=len(columns),
                                         col_widths=col_widths)

        # Column headers
        row = table.add_row()
        for i, col_def in enumerate(columns):
            cell = row.cells[i]
            ox.set_cell_shading(cell, self._label_fill)
            ox.set_cell_vertical_align(cell, col_def.vertical_align)
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(col_def.title)
            run.bold = col_def.bold
            run.font.size = self._body_size
            run.font.name = self._default_font

        # Data rows
        rows_data = section_data.get("rows", [])
        for row_data in rows_data:
            row = table.add_row()
            for i, col_def in enumerate(columns):
                cell = row.cells[i]
                ox.set_cell_vertical_align(cell, col_def.vertical_align)
                if col_def.fill_color:
                    ox.set_cell_shading(cell, col_def.fill_color)
                p = cell.paragraphs[0]
                ox.set_paragraph_spacing(p, before=60, after=60)
                value = row_data.get(col_def.id, "")
                run = p.add_run(str(value))
                run.font.size = self._body_size
                run.font.name = self._default_font

    def _build_free_text(self, section_def: SectionDefinition,
                         section_data: dict, all_data: dict):
        """Build a free text section."""
        table = self._add_section_header(section_def.title, col_count=1,
                                         col_widths=[TABLE_WIDTH])
        text = section_data.get("text", "")
        if text:
            row = table.add_row()
            cell = row.cells[0]
            p = cell.paragraphs[0]
            ox.set_paragraph_spacing(p, before=60, after=60)
            run = p.add_run(text)
            run.font.size = self._body_size
            run.font.name = self._default_font

    def _build_flowchart_section(self, section_def: SectionDefinition,
                                 section_data: dict, all_data: dict):
        """Build a flowchart section with OOXML drawing shapes."""
        from .flowchart import Flowchart
        from .flowchart_layout import FlowchartLayoutEngine

        flowchart_data = section_data.get("flowchart")
        if not flowchart_data:
            # Placeholder
            p = self.doc.add_paragraph()
            run = p.add_run(f"[Flowchart: {section_def.title}]")
            run.font.size = self._body_size
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            return

        if isinstance(flowchart_data, Flowchart):
            flowchart = flowchart_data
        else:
            layout_engine = FlowchartLayoutEngine()
            flowchart = layout_engine.layout(flowchart_data.get("steps", []))

        # Build OOXML shapes
        shapes = []
        for node in flowchart.nodes:
            if node.shape == "diamond":
                shape = ox.build_shape_diamond(
                    node.x, node.y, node.width, node.height, node.label
                )
            elif node.shape in ("oval", "ellipse"):
                shape = ox.build_shape_oval(
                    node.x, node.y, node.width, node.height, node.label
                )
            else:
                shape = ox.build_shape_rect(
                    node.x, node.y, node.width, node.height, node.label
                )
            shapes.append(shape)

        for conn in flowchart.connectors:
            # Find source and target nodes for positioning
            src = next((n for n in flowchart.nodes if n.id == conn.from_node), None)
            tgt = next((n for n in flowchart.nodes if n.id == conn.to_node), None)
            if src and tgt:
                x1 = src.x + src.width // 2
                y1 = src.y + src.height
                x2 = tgt.x + tgt.width // 2
                y2 = tgt.y
                shapes.append(ox.build_connector_shape(x1, y1, x2, y2))

        drawing = ox.build_flowchart_drawing(
            shapes, flowchart.total_width, flowchart.total_height
        )

        # Insert drawing into document
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p._p.append(drawing)
